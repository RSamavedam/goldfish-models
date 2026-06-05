"""Build the speaker-script docx.

Layout choices:
  - 11pt Calibri body, 16pt Georgia headings, 14pt slide titles
  - Each slide gets a heading + a (≈Xs) timing tag + a blockquote-styled
    speaker prose, so the reader can scan headings and read the script
    underneath.
  - Timing table at the bottom.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --------------------------------------------------------------------------
# Content
# --------------------------------------------------------------------------

SLIDES = [
    ("Slide 1 — Title", "≈25 s", """
Today I'm going to talk about Goldfish Models. The premise is simple: what if every turn a language-model agent took, it only saw the last couple thousand tokens of its own conversation history? Everything older is just gone. No retrieval, no summarization, no clever attention trick — it's just truncated.

The catch is that we then have to teach the model to use the filesystem as its memory. So: a goldfish with a notebook.
"""),
    ("Slide 2 — Biological motivation", "≈60 s", """
Before the engineering pitch, the biological intuition. The standard cognitive-science number — Cowan, 2001 — is that humans hold about four items in active working memory at once. Miller's famous 7±2 figure was for short-term capacity including chunking and rehearsal; for raw, untrained, immediate recall the number is closer to four.

That's a tiny number. And yet humans solve enormous problems. The reason isn't that we externalize everything onto paper — it's that we have an enormous long-term memory and we retrieve from it precisely. A surgeon recalls a specific protocol step. A programmer recalls a specific function signature. We're not doing fuzzy embedding-style similarity search; we're naming and pulling structured items.

And critically: when we learn a new domain, we deliberately structure that long-term store as we go. Outlines, indexes, mnemonics, hierarchies. We curate the memory so future retrieval is cheap and precise.

The goldfish proposal mirrors this. A small per-turn working set — the context window — plus a precise, structured long-term store — the filesystem. Not fuzzy embedding lookup. The model writes named files, edits them, greps them. The retrieval surface is the same shape as human long-term recall, not the same shape as RAG.
"""),
    ("Slide 3 — The KV cache problem", "≈80 s", """
So let's ground this. The KV cache for a transformer is two cache tensors — one for keys, one for values — per layer, per attention head, of dimension d_head, stored at some precision. You multiply all of that by the sequence length and by the batch size.

For a Llama-3 70B–class model at FP16, that comes out to roughly 320 kilobytes per token of context.

If you want to serve at a 32 thousand–token context window with a batch size of 16, you're looking at 164 gigabytes — for the KV cache alone, not counting weights or activations.

The implication on the right side is the important bit. The cache grows linearly with both sequence length and batch size, so doubling your context literally halves how many concurrent users you can fit on a GPU. An 80-gigabyte H100 ends up holding maybe seven simultaneous 32k-context requests — and that's nowhere near the hundreds you'd need to actually keep the FLOPs busy.

The cache is the throttle. Not compute.
"""),
    ("Slide 4 — The 15× throughput claim", "≈80 s", """
Here's the proposition concretely. Same hardware: a Llama-class 70B on an 80-gigabyte H100. Roughly 70 gigs free after weights and activations.

In the status quo on the left, each request needs around 10 gigs of KV cache at 32k tokens, which gives you about 7 concurrent users. That's the baseline.

Now suppose the agent's per-turn context window is capped at 2,048 tokens instead — what we call goldfish-2048 on the right. Per-request KV cache drops to about two-thirds of a gigabyte. With the same 70 gigs free, you can fit around 105 concurrent users.

Same hardware, same model. Roughly a 15× throughput gain on long-horizon agentic workloads. And critically, this doesn't require any change to model weights — only to the way the agent harness manages history. It's an inference-time architectural choice.
"""),
    ("Slide 5 — The proposal", "≈60 s", """
So the proposal has three pieces.

First, a hard window cap. Every turn, the model sees only the last L tokens of conversation history, where L is something like 128, 512, 1024, or 2048. We're not doing any clever summarization — we're just truncating. Naive but cheap.

Second, filesystem persistence. The agent has a private working directory between turns. It can read and write files. Its notes, its state, the repo it's editing — all of that survives across turns even though the conversation history doesn't.

Third, a light-touch protocol in the system prompt. We tell the model the regime exists — the window is small, the filesystem persists, thinking should go to disk. But we deliberately do NOT prescribe a state file format. We want to see what the model invents.
"""),
    ("Slide 6 — Prior methods", "≈85 s", """
Now there's a substantial body of prior work on KV cache compression. Let me walk through this table briefly.

H2O keeps "heavy hitter" tokens identified by online attention scores. SnapKV clusters per-attention-head. NACL trains an auxiliary saliency scorer per layer. InfiniPot keeps a continual in-place summary. Hashevict bucket-approximates with hashing. MorphKV learns a per-token retention policy. RocketKV does hierarchical multi-stage compression.

They all share two structural properties. They reduce the cache by deciding which tokens to drop, and they use the model's own internal signals — attention magnitude, hidden saliency, hash collisions — to make that decision.

That's the part we think is wrong-shaped. The assumption is that some byproduct of the forward pass tells you what's important to remember. But attention magnitude is a notoriously poor proxy for semantic importance; learned scorers don't transfer across tasks; hash collisions are not semantic similarity.

What Goldfish proposes instead is: we don't try to read this off the model. We let the model decide explicitly, in shell, what it wants to write down for next turn. That decision is inspectable and the model can be wrong about it in ways we can see.
"""),
    ("Slide 7 — Setup", "≈40 s", """
The benchmark is SWE-bench Verified — real bug-fix tasks in real Python repositories: django, astropy, sympy. The model has to produce a unified-diff patch. The official Docker-backed scorer applies the patch to the actual repo at the actual commit, runs the held-out test suite, and a patch only passes if every previously-failing test now passes and every previously-passing test still passes. So we're scoring against real test execution, not an LLM-as-judge.

On the right is why we picked gpt-4o. The honest answer is that reasoning models — gpt-5, the o-series — have hidden chain-of-thought that competes with visible output for the same token budget. We tried it. At L equals max output, every turn was zero visible tokens because the entire budget got spent on reasoning. Goldfish is just not the right regime for opaque-reasoner architectures.
"""),
    ("Slide 8 — Results", "≈70 s", """
Here's the headline. Solve rate climbs monotonically with L.

At L equals 512 — eleven cells, zero solves. That's a hard floor. Below about a thousand tokens the protocol just can't bootstrap — the model can't even fit its own notes back into the window after reading them.

At L equals 1024, we hit 18 percent. Two solves out of eleven.

At L equals 2048 — which is the interesting regime, the one that gives us the 15× throughput — we hit 25 percent.

And at L equals infinity, native context, we cap out at 33 percent.

So the L=2048 condition reaches roughly 75 percent of the native capability with a 16× smaller per-turn context. That's the headline finding: the protocol is doing real work, not just preserving a token-budget property.

One thing the bar chart doesn't show: the gap is task-specific. Some tasks solve at every L from 1024 up. One task is native-only. Most tasks don't solve at any L — that's gpt-4o's capability ceiling, not the goldfish regime's.
"""),
    ("Slide 9 — Emergent behaviors", "≈70 s", """
The interesting part — what does the model actually do when L is small. None of these behaviors were specified by the prompt.

First, the model uses shell comments as inline notes. Instead of writing a separate notes.md file like we suggested, it prefixes each command with hash-comments. Bash ignores them, but they appear twice in history.txt — once as input, once as command echo — which makes them survive truncation more robustly than a separate file does. The model figured out that the cheapest persistent memory in this regime is a comment line in its own commands.

Second, the model uses prose between code blocks as scratch memory. On what I started calling "hypothesis turns," typically turn three of a solve, the model writes two or three sentences diagnosing the bug. That prose primes the next several turns. Visible output is doing double duty as both reasoning and persistent state.

Third, every single solving trajectory follows the same tool-fallback pattern: try sed, fail because of escaped characters, fall back to a python3 heredoc. The model learned this preference on the fly, across three different L values and three different tasks.

And fourth, the model self-compresses. Average response length goes DOWN as L goes down — about 68 tokens per turn at L=1024 versus 92 at L equals infinity. The constraint is actually a productivity feature, not just a memory limit.
"""),
    ("Slide 10 — Takeaways", "≈40 s", """
To wrap up.

The hard problem in KV cache compression isn't engineering the compression — it's deciding what to keep. Prior work answers this by reading model internals; Goldfish externalizes the decision into the agent's filesystem.

Most of native-context capability survives a 16× context cut — 75 percent solve rate at L=2048 with a roughly 15× smaller KV footprint. The protocol is doing real work.

There's a floor. Below about a thousand tokens, the protocol can't bootstrap. The act of reading your own notes consumes the window.

And the regime is for non-reasoning models. Reasoning architectures like gpt-5 collapse because hidden chain-of-thought competes with visible output for the same budget.

Code's on GitHub. Thank you.
"""),
]

TIMING_TABLE = [
    ("Slide", "Approx s", "Cum"),
    ("1",  "25", "0:25"),
    ("2",  "55", "1:20"),
    ("3",  "80", "2:40"),
    ("4",  "80", "4:00"),
    ("5",  "60", "5:00"),
    ("6",  "85", "6:25"),
    ("7",  "40", "7:05"),
    ("8",  "70", "8:15"),
    ("9",  "70", "9:25"),
    ("10", "40", "10:05"),
]

# --------------------------------------------------------------------------
# Build
# --------------------------------------------------------------------------

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Defaults
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("Goldfish Models — Speaker Script")
run.font.name = "Georgia"
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x21, 0x29, 0x5C)

subtitle = doc.add_paragraph()
sub_run = subtitle.add_run("Target: ≈9 minutes total · pace ≈150 wpm · hit the numbers slowly")
sub_run.font.size = Pt(11)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(0x5A, 0x6A, 0x7F)

# Spacer
doc.add_paragraph()

# --- slides ---
for heading, timing, body in SLIDES:
    # Heading line
    h = doc.add_paragraph()
    h_run = h.add_run(heading)
    h_run.font.name = "Georgia"
    h_run.font.size = Pt(15)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(0x06, 0x5A, 0x82)

    # Inline timing
    timing_run = h.add_run(f"    ({timing})")
    timing_run.font.name = "Calibri"
    timing_run.font.size = Pt(11)
    timing_run.font.italic = True
    timing_run.font.color.rgb = RGBColor(0xFF, 0x7F, 0x50)

    # Speaker text — each paragraph in body, indented like a blockquote
    paras = [p.strip() for p in body.strip().split("\n\n") if p.strip()]
    for ptext in paras:
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Inches(0.3)
        bp.paragraph_format.space_after = Pt(6)
        # Subtle left-bar effect via a single em quad before the text? Skip;
        # rely on indent + italic-light styling for the script voice.
        bp_run = bp.add_run(ptext)
        bp_run.font.name = "Calibri"
        bp_run.font.size = Pt(11)
        bp_run.font.color.rgb = RGBColor(0x1E, 0x27, 0x61)

    # Small gap between slides
    doc.add_paragraph()

# --- Timing table ---
h = doc.add_paragraph()
hr = h.add_run("Timing summary")
hr.font.name = "Georgia"
hr.font.size = Pt(15)
hr.font.bold = True
hr.font.color.rgb = RGBColor(0x06, 0x5A, 0x82)

table = doc.add_table(rows=len(TIMING_TABLE), cols=3)
table.style = "Light Grid"
for r_idx, row_vals in enumerate(TIMING_TABLE):
    cells = table.rows[r_idx].cells
    for c_idx, val in enumerate(row_vals):
        cells[c_idx].text = val
        for paragraph in cells[c_idx].paragraphs:
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                if r_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                else:
                    run.font.color.rgb = RGBColor(0x1E, 0x27, 0x61)
    if r_idx == 0:
        # Shade header row by writing into XML (python-docx doesn't expose
        # cell shading cleanly).
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        for cell in cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "21295C")
            tc_pr.append(shd)

doc.add_paragraph()

# --- Contingencies ---
h = doc.add_paragraph()
hr = h.add_run("If you're running long")
hr.font.name = "Georgia"
hr.font.size = Pt(13)
hr.font.bold = True
hr.font.color.rgb = RGBColor(0x06, 0x5A, 0x82)

for line in [
    "Slide 6 can lose the per-method walkthrough — just say \"all of these compress by deciding which tokens to drop using a model-internal signal\" and point at the table.",
    "Slide 9 can drop the self-compression bullet.",
    "Slide 2 (bio motivation) can be cut to just the 4±1 stat and the closing line: \"LLMs hold everything in context; we're making them act more like us.\" Saves ~25 seconds.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p_run = p.add_run(line)
    p_run.font.name = "Calibri"
    p_run.font.size = Pt(11)

h = doc.add_paragraph()
hr = h.add_run("If you're running short")
hr.font.name = "Georgia"
hr.font.size = Pt(13)
hr.font.bold = True
hr.font.color.rgb = RGBColor(0x06, 0x5A, 0x82)
p = doc.add_paragraph(style="List Bullet")
p_run = p.add_run(
    "Add to Slide 8: \"and on the L=2048 tasks that DO solve, average wall-clock is "
    "9 turns — versus 10 at native. The compression doesn't slow trajectories down.\""
)
p_run.font.name = "Calibri"
p_run.font.size = Pt(11)

out = "/Users/raghavsamavedam/Documents/153_project/goldfish-models/paper/talk/script.docx"
doc.save(out)
print(f"wrote {out}")
